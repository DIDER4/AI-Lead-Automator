"""
Knowledge Base Service
Handles document ingestion, embedding, and semantic search using RAG
"""

import os
import uuid
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import json

# Document processing
import PyPDF2
import docx

# Embeddings and vector storage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangChainDocument

from src.config import get_logger, Constants
from src.models.document import Document

logger = get_logger(__name__)


class KnowledgeBaseError(Exception):
    """Custom exception for knowledge base errors"""
    pass


class KnowledgeBaseService:
    """
    RAG-based knowledge base service
    Handles document ingestion, chunking, embedding, and semantic search
    """
    
    def __init__(self, persist_directory: str = "data/chroma_db", 
                 documents_directory: str = "data/documents"):
        """
        Initialize Knowledge Base Service
        
        Args:
            persist_directory: Directory for ChromaDB persistence
            documents_directory: Directory to store original uploaded files
        """
        self.persist_directory = Path(persist_directory)
        self.documents_directory = Path(documents_directory)
        self.metadata_file = self.documents_directory / "documents_metadata.json"
        
        # Create directories if they don't exist
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        self.documents_directory.mkdir(parents=True, exist_ok=True)
        
        # Initialize embeddings model (local, no API calls)
        logger.info("Initializing embeddings model (this may take a moment on first run)...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",  # Fast and efficient
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize or load vector store
        self.vector_store = self._initialize_vector_store()
        
        logger.info("Knowledge Base Service initialized")
    
    def _initialize_vector_store(self) -> Chroma:
        """Initialize or load existing ChromaDB vector store"""
        try:
            # Try to load existing collection
            vector_store = Chroma(
                persist_directory=str(self.persist_directory),
                embedding_function=self.embeddings,
                collection_name="company_knowledge"
            )
            logger.info(f"Loaded existing vector store with {vector_store._collection.count()} documents")
            return vector_store
        except Exception as e:
            logger.warning(f"Could not load existing vector store: {e}")
            # Create new collection
            vector_store = Chroma(
                persist_directory=str(self.persist_directory),
                embedding_function=self.embeddings,
                collection_name="company_knowledge"
            )
            logger.info("Created new vector store")
            return vector_store
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise KnowledgeBaseError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise KnowledgeBaseError(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise KnowledgeBaseError(f"Failed to read text file: {str(e)}")
    
    def add_document(self, file_path: str, filename: str) -> Tuple[bool, str, Optional[Document]]:
        """
        Add document to knowledge base
        
        Args:
            file_path: Path to uploaded file (temporary)
            filename: Original filename
            
        Returns:
            Tuple of (success, message, document)
        """
        try:
            file_path = Path(file_path)
            
            # Determine document type
            file_extension = file_path.suffix.lower()
            if file_extension == '.pdf':
                doc_type = 'pdf'
                text = self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                doc_type = 'docx'
                text = self._extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                doc_type = 'txt'
                text = self._extract_text_from_txt(file_path)
            else:
                return False, f"Unsupported file type: {file_extension}", None
            
            if not text or len(text) < 50:
                return False, "Document appears to be empty or too short", None
            
            # Generate unique ID
            doc_id = str(uuid.uuid4())
            
            # Save original file
            saved_file_path = self.documents_directory / f"{doc_id}_{filename}"
            if file_path != saved_file_path:
                import shutil
                shutil.copy2(file_path, saved_file_path)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Split document into {len(chunks)} chunks")
            
            # Create LangChain documents with metadata
            langchain_docs = []
            for i, chunk in enumerate(chunks):
                langchain_docs.append(LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "doc_id": doc_id,
                        "chunk_index": i,
                        "doc_type": doc_type
                    }
                ))
            
            # Add to vector store
            self.vector_store.add_documents(langchain_docs)
            # Note: langchain-chroma persists automatically
            
            # Calculate metrics
            char_count = len(text)
            token_count = Document.estimate_tokens(text)
            avg_chunk_size = char_count / len(chunks) if chunks else 0
            
            # Estimate embedding cost (OpenAI ada-002: $0.0001 per 1K tokens)
            # We use local embeddings (free), but estimate if user switches
            embedding_cost = (token_count / 1000) * 0.0001
            
            # Get file modification time
            last_modified = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            
            # Generate summary (first 200 chars)
            summary = text[:200].replace('\n', ' ').strip()
            if len(text) > 200:
                summary += "..."
            
            # Create Document object
            document = Document(
                filename=filename,
                content=text[:1000] + "..." if len(text) > 1000 else text,  # Store preview
                doc_type=doc_type,
                file_size=file_path.stat().st_size,
                num_chunks=len(chunks),
                file_path=str(saved_file_path),
                id=doc_id,
                char_count=char_count,
                token_count=token_count,
                avg_chunk_size=round(avg_chunk_size, 1),
                last_modified=last_modified,
                embedding_cost_estimate=round(embedding_cost, 6),
                summary=summary
            )
            
            # Save metadata
            self._save_document_metadata(document)
            
            logger.info(f"Successfully added document: {filename} ({len(chunks)} chunks)")
            return True, f"Document '{filename}' successfully added to knowledge base", document
            
        except Exception as e:
            logger.error(f"Error adding document: {e}", exc_info=True)
            return False, f"Error adding document: {str(e)}", None
    
    def search(self, query: str, k: int = 5) -> List[Dict]:
        """
        Perform semantic search in knowledge base
        
        Args:
            query: Search query
            k: Number of results to return
            
        Returns:
            List of relevant document chunks with metadata
        """
        try:
            if self.vector_store._collection.count() == 0:
                logger.info("Vector store is empty, returning no results")
                return []
            
            # Perform similarity search
            results = self.vector_store.similarity_search_with_score(query, k=k)
            
            # Format results
            formatted_results = []
            for doc, score in results:
                formatted_results.append({
                    "content": doc.page_content,
                    "score": float(score),
                    "metadata": doc.metadata
                })
            
            logger.info(f"Search returned {len(formatted_results)} results")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error during search: {e}", exc_info=True)
            return []
    
    def get_context_for_prompt(self, query: str, max_chunks: int = 3) -> str:
        """
        Get formatted context for AI prompt
        
        Args:
            query: Search query
            max_chunks: Maximum number of chunks to include
            
        Returns:
            Formatted context string
        """
        results = self.search(query, k=max_chunks)
        
        if not results:
            return ""
        
        context_parts = []
        for i, result in enumerate(results, 1):
            source = result['metadata'].get('source', 'Unknown')
            content = result['content']
            context_parts.append(f"[Source {i}: {source}]\n{content}")
        
        context = "\n\n".join(context_parts)
        return context
    
    def list_documents(self) -> List[Document]:
        """
        List all documents in knowledge base
        
        Returns:
            List of Document objects
        """
        try:
            if not self.metadata_file.exists():
                return []
            
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                metadata_list = json.load(f)
            
            documents = [Document.from_dict(meta) for meta in metadata_list]
            return documents
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []
    
    def delete_document(self, doc_id: str) -> Tuple[bool, str]:
        """
        Delete document from knowledge base
        
        Args:
            doc_id: Document ID
            
        Returns:
            Tuple of (success, message)
        """
        try:
            # Delete from vector store
            self.vector_store.delete(where={"doc_id": doc_id})
            # Note: langchain-chroma persists automatically
            
            # Delete file
            documents = self.list_documents()
            doc_to_delete = None
            for doc in documents:
                if doc.id == doc_id:
                    doc_to_delete = doc
                    if doc.file_path and Path(doc.file_path).exists():
                        Path(doc.file_path).unlink()
                    break
            
            if doc_to_delete:
                # Update metadata
                documents.remove(doc_to_delete)
                self._save_all_documents_metadata(documents)
                
                logger.info(f"Deleted document: {doc_id}")
                return True, "Document deleted successfully"
            else:
                return False, "Document not found"
                
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False, f"Error deleting document: {str(e)}"
    
    def get_stats(self) -> Dict:
        """Get knowledge base statistics with enhanced metrics"""
        documents = self.list_documents()
        
        total_chunks = self.vector_store._collection.count() if self.vector_store else 0
        
        # Calculate aggregate metrics
        total_chars = sum(doc.char_count for doc in documents)
        total_tokens = sum(doc.token_count for doc in documents)
        total_cost = sum(doc.embedding_cost_estimate for doc in documents)
        avg_doc_size = total_chars / len(documents) if documents else 0
        
        return {
            "total_documents": len(documents),
            "total_chunks": total_chunks,
            "doc_types": self._count_doc_types(documents),
            "total_characters": total_chars,
            "total_tokens": total_tokens,
            "total_embedding_cost": round(total_cost, 4),
            "avg_doc_size": round(avg_doc_size, 0)
        }
    
    def _count_doc_types(self, documents: List[Document]) -> Dict[str, int]:
        """Count documents by type"""
        counts = {}
        for doc in documents:
            counts[doc.doc_type] = counts.get(doc.doc_type, 0) + 1
        return counts
    
    def _save_document_metadata(self, document: Document):
        """Save single document metadata"""
        documents = self.list_documents()
        documents.append(document)
        self._save_all_documents_metadata(documents)
    
    def _save_all_documents_metadata(self, documents: List[Document]):
        """Save all documents metadata"""
        try:
            metadata_list = [doc.to_dict() for doc in documents]
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata_list, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
